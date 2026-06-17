"""
ResumeUploadAgent

Responsibilities:
  - Validates file type (PDF, DOCX, TXT)
  - Extracts raw text using pdfplumber (PDF) or python-docx (DOCX)
  - Chunks text into overlapping 512-token segments (50-token overlap)
  - Stores file metadata in state
  - NEVER hallucinate content from images or embedded objects
"""

from __future__ import annotations

import logging
import math
import os
import re
from typing import Optional

from app.agents.base_agent import BaseAgent
from app.agents.state import AgentState

logger = logging.getLogger(__name__)

_SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
_CHUNK_TOKENS = 512
_OVERLAP_TOKENS = 50
# Approximate chars-per-token for English text
_CHARS_PER_TOKEN = 4


class ResumeUploadAgent(BaseAgent):
    """
    First agent in the resume-analysis pipeline.
    Extracts raw text and chunked segments from the uploaded file.
    """

    async def run(self, state: AgentState) -> AgentState:
        try:
            file_path: Optional[str] = state.get("resume_file_path")
            if not file_path:
                return self._log_error(state, "No resume_file_path provided in state")

            if not os.path.exists(file_path):
                return self._log_error(state, f"File not found: {file_path}")

            ext = os.path.splitext(file_path)[1].lower()
            if ext not in _SUPPORTED_EXTENSIONS:
                return self._log_error(
                    state,
                    f"Unsupported file type '{ext}'. Accepted: {_SUPPORTED_EXTENSIONS}",
                )

            # ---------------------------------------------------------------- #
            # Extract raw text
            # ---------------------------------------------------------------- #
            raw_text, extraction_quality = await self._extract_text(file_path, ext)

            if not raw_text or len(raw_text.strip()) < 50:
                state = self._log_error(
                    state,
                    "Extracted text is too short or empty — resume may be image-only. "
                    "Please provide a text-based PDF or DOCX.",
                )
                state = self._update_confidence(state, 0.0)
                return state

            # ---------------------------------------------------------------- #
            # Chunk the text
            # ---------------------------------------------------------------- #
            chunks = self._chunk_text(raw_text, _CHUNK_TOKENS, _OVERLAP_TOKENS)

            # ---------------------------------------------------------------- #
            # Write to state
            # ---------------------------------------------------------------- #
            state["resume_raw_text"] = raw_text
            state["resume_chunks"] = chunks

            # Confidence score based on extraction quality and text length
            confidence = self._compute_confidence(raw_text, extraction_quality)
            state = self._update_confidence(state, confidence)
            state = self._append_message(
                state,
                role="system",
                content=(
                    f"ResumeUploadAgent: Extracted {len(raw_text)} chars, "
                    f"{len(chunks)} chunks from {os.path.basename(file_path)}. "
                    f"Confidence={confidence:.2f}"
                ),
            )
            logger.info(
                "[ResumeUploadAgent] Extracted %d chars, %d chunks from %s",
                len(raw_text), len(chunks), file_path,
            )

        except Exception as exc:
            state = self._log_error(state, f"Unexpected error: {exc}")
            state = self._update_confidence(state, 0.0)

        return state

    # ---------------------------------------------------------------------- #
    # Private helpers
    # ---------------------------------------------------------------------- #

    async def _extract_text(self, file_path: str, ext: str) -> tuple[str, float]:
        """
        Extract text and return (text, quality_score 0-1).
        quality_score reflects how much structured text was recovered.
        """
        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext in {".docx", ".doc"}:
            return self._extract_docx(file_path)
        else:  # .txt
            return self._extract_txt(file_path)

    def _extract_pdf(self, file_path: str) -> tuple[str, float]:
        try:
            import pdfplumber  # type: ignore
        except ImportError:
            raise ImportError("pdfplumber required: pip install pdfplumber")

        pages_text: list[str] = []
        total_pages = 0
        pages_with_text = 0

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(text)
                    pages_with_text += 1

        raw = "\n\n".join(pages_text)
        quality = pages_with_text / max(total_pages, 1)
        return raw, quality

    def _extract_docx(self, file_path: str) -> tuple[str, float]:
        try:
            from docx import Document  # type: ignore
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        raw = "\n".join(paragraphs)
        quality = 1.0 if paragraphs else 0.0
        return raw, quality

    def _extract_txt(self, file_path: str) -> tuple[str, float]:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        quality = 1.0 if raw.strip() else 0.0
        return raw, quality

    def _chunk_text(self, text: str, chunk_tokens: int, overlap_tokens: int) -> list[str]:
        """
        Split text into overlapping chunks.
        Uses character-based approximation (4 chars ≈ 1 token).
        """
        chunk_chars = chunk_tokens * _CHARS_PER_TOKEN
        overlap_chars = overlap_tokens * _CHARS_PER_TOKEN
        step = chunk_chars - overlap_chars

        # Split on sentence boundaries when possible
        sentences = re.split(r"(?<=[.!?])\s+|\n{2,}", text)
        chunks: list[str] = []
        current = ""

        for sentence in sentences:
            if len(current) + len(sentence) + 1 <= chunk_chars:
                current = (current + " " + sentence).strip()
            else:
                if current:
                    chunks.append(current)
                # Start next chunk with overlap
                overlap_text = current[-overlap_chars:] if len(current) > overlap_chars else current
                current = (overlap_text + " " + sentence).strip()

        if current:
            chunks.append(current)

        # Fallback: if single block without sentence breaks
        if not chunks:
            for i in range(0, len(text), step):
                chunk = text[i: i + chunk_chars]
                if chunk.strip():
                    chunks.append(chunk.strip())

        return chunks

    def _compute_confidence(self, text: str, extraction_quality: float) -> float:
        """
        Heuristic confidence based on text length and extraction quality.
        """
        length_score = min(len(text) / 3000, 1.0)  # saturate at 3000 chars
        # Check for key resume sections
        section_hits = sum(
            1
            for kw in ["experience", "education", "skills", "summary", "objective"]
            if kw in text.lower()
        )
        section_score = min(section_hits / 3, 1.0)

        confidence = (extraction_quality * 0.4) + (length_score * 0.35) + (section_score * 0.25)
        return round(confidence, 4)
